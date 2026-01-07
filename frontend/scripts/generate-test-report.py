#!/usr/bin/env python3
"""
Generate DOCX Test Report for Unit Tests
Tests for: Sneha Iyer (MANAGER), Ravi Krishnan (MANAGER), Divya Menon (EMPLOYEE)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os

def create_test_report():
    doc = Document()

    # Title
    title = doc.add_heading('Unit Test Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Zenora HR Platform - API Testing')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run('Test Status: ').bold = True
    pass_run = summary.add_run('ALL TESTS PASSED ✓')
    pass_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    pass_run.bold = True

    doc.add_paragraph()

    # Summary Table
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Table Grid'

    summary_data = [
        ('Total Test Suites', '3'),
        ('Total Tests', '46'),
        ('Passed', '46'),
        ('Failed', '0'),
        ('Execution Time', '~3.0 seconds'),
    ]

    for i, (label, value) in enumerate(summary_data):
        row = summary_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Test Users Section
    doc.add_heading('Test Users', level=1)
    doc.add_paragraph('The following users were tested across all API endpoints:')

    users_table = doc.add_table(rows=4, cols=4)
    users_table.style = 'Table Grid'

    # Header
    header_cells = users_table.rows[0].cells
    headers = ['User Name', 'Email', 'Role', 'Tenant']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True

    users_data = [
        ('Sneha Iyer', 'sneha.iyer@idwteam.com', 'MANAGER', 'tenant-idwteam-001'),
        ('Ravi Krishnan', 'ravi.krishnan@idwteam.com', 'MANAGER', 'tenant-idwteam-001'),
        ('Divya Menon', 'divya.menon@idwteam.com', 'EMPLOYEE', 'tenant-idwteam-001'),
    ]

    for i, (name, email, role, tenant) in enumerate(users_data):
        row = users_table.rows[i + 1]
        row.cells[0].text = name
        row.cells[1].text = email
        row.cells[2].text = role
        row.cells[3].text = tenant

    doc.add_paragraph()

    # Test Suite 1: Employee Timesheets
    doc.add_heading('Test Suite 1: Employee Timesheets API', level=1)
    doc.add_paragraph('File: __tests__/api/employee/timesheets.test.ts')

    timesheet_tests = [
        ('Authentication Tests', [
            ('should return 401 when no session cookie is provided', 'PASS'),
            ('should return 401 for invalid session', 'PASS'),
        ]),
        ('Timesheet Fetch Tests for Sneha Iyer (MANAGER)', [
            ('should fetch timesheets for Sneha Iyer successfully', 'PASS'),
            ('should return 400 when startDate or endDate is missing', 'PASS'),
        ]),
        ('Timesheet Fetch Tests for Ravi Krishnan (MANAGER)', [
            ('should fetch timesheets for Ravi Krishnan successfully', 'PASS'),
            ('should handle empty timesheet result for Ravi', 'PASS'),
        ]),
        ('Timesheet Fetch Tests for Divya Menon (EMPLOYEE)', [
            ('should fetch timesheets for Divya Menon (EMPLOYEE) successfully', 'PASS'),
        ]),
        ('Pagination Tests', [
            ('should return 400 for invalid pagination parameters', 'PASS'),
            ('should handle pagination correctly', 'PASS'),
        ]),
        ('Create Timesheet Tests for Divya Menon (EMPLOYEE)', [
            ('should create a new timesheet entry for Divya', 'PASS'),
            ('should return 400 when required fields are missing', 'PASS'),
            ('should return 400 for invalid hoursWorked value', 'PASS'),
        ]),
        ('Create Timesheet Tests for Managers', [
            ('should create timesheet for Sneha Iyer (MANAGER)', 'PASS'),
        ]),
    ]

    for section_name, tests in timesheet_tests:
        doc.add_heading(section_name, level=2)
        table = doc.add_table(rows=len(tests) + 1, cols=2)
        table.style = 'Table Grid'

        header_row = table.rows[0]
        header_row.cells[0].text = 'Test Case'
        header_row.cells[1].text = 'Status'
        header_row.cells[0].paragraphs[0].runs[0].bold = True
        header_row.cells[1].paragraphs[0].runs[0].bold = True

        for i, (test_name, status) in enumerate(tests):
            row = table.rows[i + 1]
            row.cells[0].text = test_name
            row.cells[1].text = status
            if status == 'PASS':
                row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        doc.add_paragraph()

    # Test Suite 2: Employee Expenses
    doc.add_heading('Test Suite 2: Employee Expenses API', level=1)
    doc.add_paragraph('File: __tests__/api/employee/expenses.test.ts')

    expense_tests = [
        ('Authentication Tests', [
            ('should return 401 when no session cookie is provided', 'PASS'),
            ('should return 401 for invalid session', 'PASS'),
        ]),
        ('Expense Fetch Tests for Sneha Iyer (MANAGER)', [
            ('should fetch expenses for Sneha Iyer successfully', 'PASS'),
            ('should filter expenses by date range', 'PASS'),
        ]),
        ('Expense Fetch Tests for Ravi Krishnan (MANAGER)', [
            ('should fetch expenses for Ravi Krishnan successfully', 'PASS'),
            ('should filter expenses by status', 'PASS'),
        ]),
        ('Expense Fetch Tests for Divya Menon (EMPLOYEE)', [
            ('should fetch expenses for Divya Menon (EMPLOYEE) successfully', 'PASS'),
        ]),
        ('Create Expense Tests for Divya Menon (EMPLOYEE)', [
            ('should create a new expense claim for Divya', 'PASS'),
            ('should return 400 when required fields are missing', 'PASS'),
            ('should return 400 for amount less than or equal to 0', 'PASS'),
            ('should return 400 for short description', 'PASS'),
            ('should return 400 for expense over $50 without receipt', 'PASS'),
            ('should return 400 for future expense date', 'PASS'),
            ('should return 409 for duplicate expense', 'PASS'),
        ]),
        ('Create Expense Tests for Managers', [
            ('should create expense for Sneha Iyer (MANAGER)', 'PASS'),
            ('should create expense for Ravi Krishnan (MANAGER)', 'PASS'),
        ]),
    ]

    for section_name, tests in expense_tests:
        doc.add_heading(section_name, level=2)
        table = doc.add_table(rows=len(tests) + 1, cols=2)
        table.style = 'Table Grid'

        header_row = table.rows[0]
        header_row.cells[0].text = 'Test Case'
        header_row.cells[1].text = 'Status'
        header_row.cells[0].paragraphs[0].runs[0].bold = True
        header_row.cells[1].paragraphs[0].runs[0].bold = True

        for i, (test_name, status) in enumerate(tests):
            row = table.rows[i + 1]
            row.cells[0].text = test_name
            row.cells[1].text = status
            if status == 'PASS':
                row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        doc.add_paragraph()

    # Test Suite 3: Manager Approvals
    doc.add_heading('Test Suite 3: Manager Approvals API', level=1)
    doc.add_paragraph('File: __tests__/api/manager/approvals.test.ts')

    approval_tests = [
        ('Timesheet Approvals - Authentication and Authorization', [
            ('should return 401 when no session cookie is provided', 'PASS'),
            ('should return 403 for EMPLOYEE role trying to access approvals', 'PASS'),
        ]),
        ('Timesheet Approvals for Sneha Iyer (MANAGER)', [
            ('should fetch pending timesheets for Sneha direct reports', 'PASS'),
            ('should return empty result when no direct reports', 'PASS'),
            ('should filter by date range', 'PASS'),
        ]),
        ('Timesheet Approvals for Ravi Krishnan (MANAGER)', [
            ('should return empty when Ravi has no direct reports', 'PASS'),
            ('should handle search functionality', 'PASS'),
        ]),
        ('Pagination Tests', [
            ('should return 400 for invalid pagination parameters', 'PASS'),
            ('should handle pagination correctly', 'PASS'),
        ]),
        ('Expense Approvals - Authentication and Authorization', [
            ('should return 401 when no session cookie is provided', 'PASS'),
            ('should return 403 for EMPLOYEE role', 'PASS'),
        ]),
        ('Expense Approvals for Sneha Iyer (MANAGER)', [
            ('should fetch pending expenses for Sneha direct reports (Divya)', 'PASS'),
            ('should return empty when no direct reports have pending expenses', 'PASS'),
            ('should filter by date range', 'PASS'),
        ]),
        ('Expense Approvals for Ravi Krishnan (MANAGER)', [
            ('should return empty when Ravi has no direct reports', 'PASS'),
        ]),
        ('Cross-User Approval Scenarios', [
            ('should show Divya timesheets in Sneha pending queue', 'PASS'),
            ('should show Divya expenses in Sneha pending queue', 'PASS'),
        ]),
    ]

    for section_name, tests in approval_tests:
        doc.add_heading(section_name, level=2)
        table = doc.add_table(rows=len(tests) + 1, cols=2)
        table.style = 'Table Grid'

        header_row = table.rows[0]
        header_row.cells[0].text = 'Test Case'
        header_row.cells[1].text = 'Status'
        header_row.cells[0].paragraphs[0].runs[0].bold = True
        header_row.cells[1].paragraphs[0].runs[0].bold = True

        for i, (test_name, status) in enumerate(tests):
            row = table.rows[i + 1]
            row.cells[0].text = test_name
            row.cells[1].text = status
            if status == 'PASS':
                row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        doc.add_paragraph()

    # Test Coverage Summary
    doc.add_heading('Test Coverage Summary', level=1)

    coverage_table = doc.add_table(rows=5, cols=3)
    coverage_table.style = 'Table Grid'

    coverage_header = coverage_table.rows[0].cells
    coverage_header[0].text = 'API Endpoint'
    coverage_header[1].text = 'Tests'
    coverage_header[2].text = 'Coverage Areas'
    for cell in coverage_header:
        cell.paragraphs[0].runs[0].bold = True

    coverage_data = [
        ('GET /api/employee/timesheets', '9', 'Auth, Fetch, Pagination, Filtering'),
        ('POST /api/employee/timesheets', '4', 'Creation, Validation, Authorization'),
        ('GET /api/employee/expenses', '7', 'Auth, Fetch, Filtering, Status'),
        ('POST /api/employee/expenses', '9', 'Creation, Validation, Duplicates, Receipts'),
        ('GET /api/manager/timesheets/pending', '9', 'Auth, Authorization, Filtering, Pagination'),
        ('GET /api/manager/expenses/pending', '8', 'Auth, Authorization, Direct Reports'),
    ]

    # Need to add more rows
    for i in range(len(coverage_data) - len(coverage_table.rows) + 1):
        coverage_table.add_row()

    for i, (endpoint, tests, areas) in enumerate(coverage_data):
        row = coverage_table.rows[i + 1]
        row.cells[0].text = endpoint
        row.cells[1].text = tests
        row.cells[2].text = areas

    doc.add_paragraph()

    # Conclusion
    doc.add_heading('Conclusion', level=1)
    conclusion = doc.add_paragraph()
    conclusion.add_run('All 46 unit tests have passed successfully. ').bold = True
    conclusion.add_run(
        'The API endpoints for timesheets, expenses, and approvals are functioning correctly '
        'for all three test users: Sneha Iyer (MANAGER), Ravi Krishnan (MANAGER), and Divya Menon (EMPLOYEE). '
        'The tests cover authentication, authorization, validation, pagination, and cross-user approval workflows.'
    )

    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('— End of Report —')
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save the document
    output_path = '/Volumes/E/zenora/frontend/test-report.docx'
    doc.save(output_path)
    print(f'Report generated: {output_path}')
    return output_path

if __name__ == '__main__':
    create_test_report()
