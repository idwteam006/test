#!/usr/bin/env python3
"""
Convert Markdown Documentation to DOCX Format

This script converts the Zenora documentation markdown files to Word (.docx) format.
It uses pypandoc as the primary method and falls back to python-docx if needed.

Usage:
    python3 convert_to_docx.py

Requires:
    - pypandoc (pip install pypandoc)
    - pandoc (brew install pandoc on macOS, or apt-get install pandoc on Linux)

Alternative (if pandoc not available):
    - python-docx (pip install python-docx)
    - markdown (pip install markdown)
"""

import os
import sys
from pathlib import Path

def check_dependencies():
    """Check if required dependencies are available."""
    try:
        import pypandoc
        print("✅ pypandoc is installed")
        # Check if pandoc executable is available
        try:
            pypandoc.get_pandoc_version()
            print(f"✅ pandoc version: {pypandoc.get_pandoc_version()}")
            return 'pypandoc'
        except Exception as e:
            print(f"❌ pandoc executable not found: {e}")
            print("   Install with: brew install pandoc (macOS) or apt-get install pandoc (Linux)")
    except ImportError:
        print("❌ pypandoc not installed")
        print("   Install with: pip install pypandoc")

    # Check fallback option
    try:
        import docx
        import markdown
        print("✅ python-docx and markdown are installed (fallback method available)")
        return 'docx'
    except ImportError:
        print("❌ Fallback method not available")
        print("   Install with: pip install python-docx markdown")
        return None

def convert_with_pypandoc(input_file, output_file):
    """Convert markdown to docx using pypandoc."""
    import pypandoc

    print(f"\n📄 Converting {input_file} to {output_file}...")

    extra_args = [
        '--toc',  # Table of contents
        '--toc-depth=3',  # TOC depth
        '--highlight-style=tango',  # Code highlighting
        '--reference-doc=/dev/null',  # Use default Word template
    ]

    try:
        pypandoc.convert_file(
            input_file,
            'docx',
            outputfile=output_file,
            extra_args=extra_args
        )
        print(f"✅ Successfully converted to {output_file}")
        return True
    except Exception as e:
        print(f"❌ Conversion failed: {e}")
        return False

def convert_with_docx(input_file, output_file):
    """Convert markdown to docx using python-docx (fallback method)."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import markdown
    import re

    print(f"\n📄 Converting {input_file} to {output_file} (fallback method)...")
    print("⚠️  Note: Fallback method has limited formatting support")

    # Read markdown file
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Create Word document
    doc = Document()

    # Parse markdown line by line for better control
    lines = md_content.split('\n')
    in_code_block = False
    code_block_content = []

    for line in lines:
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block - add accumulated content
                code_text = '\n'.join(code_block_content)
                p = doc.add_paragraph(code_text)
                p.style = 'Intense Quote'
                code_block_content = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue

        if in_code_block:
            code_block_content.append(line)
            continue

        # Handle headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)

        # Handle horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)

        # Handle lists
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Remove markdown bold/italic
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')

        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            doc.add_paragraph(text, style='List Number')

        # Handle tables (basic support)
        elif '|' in line and line.strip().startswith('|'):
            # Skip table header separator lines
            if re.match(r'^\|[\s\-:]+\|', line):
                continue
            # Basic table row - just add as paragraph for now
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            doc.add_paragraph(' | '.join(cells))

        # Handle regular paragraphs
        elif line.strip():
            text = line
            # Remove markdown bold/italic
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            # Remove markdown links but keep text
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
            doc.add_paragraph(text)

    # Save document
    try:
        doc.save(output_file)
        print(f"✅ Successfully converted to {output_file}")
        return True
    except Exception as e:
        print(f"❌ Conversion failed: {e}")
        return False

def main():
    """Main conversion function."""
    print("=" * 60)
    print("Zenora Documentation - Markdown to DOCX Converter")
    print("=" * 60)

    # Check dependencies
    converter = check_dependencies()

    if not converter:
        print("\n❌ No conversion method available. Please install dependencies:")
        print("   Method 1 (recommended): pip install pypandoc && brew install pandoc")
        print("   Method 2 (fallback): pip install python-docx markdown")
        sys.exit(1)

    # Define files to convert
    files_to_convert = [
        ('ZENORA_COMPLETE_FLOWS.md', 'ZENORA_COMPLETE_FLOWS.docx'),
        ('ZENORA_COMPLETE_FEATURES.md', 'ZENORA_COMPLETE_FEATURES.docx'),
    ]

    success_count = 0
    failed_count = 0

    # Convert each file
    for input_file, output_file in files_to_convert:
        if not os.path.exists(input_file):
            print(f"\n⚠️  Skipping {input_file} - file not found")
            continue

        if converter == 'pypandoc':
            result = convert_with_pypandoc(input_file, output_file)
        else:
            result = convert_with_docx(input_file, output_file)

        if result:
            success_count += 1
            # Get file size
            size_mb = os.path.getsize(output_file) / (1024 * 1024)
            print(f"   📊 File size: {size_mb:.2f} MB")
        else:
            failed_count += 1

    # Summary
    print("\n" + "=" * 60)
    print("Conversion Summary")
    print("=" * 60)
    print(f"✅ Successful conversions: {success_count}")
    print(f"❌ Failed conversions: {failed_count}")

    if success_count > 0:
        print("\n📁 Generated files:")
        for _, output_file in files_to_convert:
            if os.path.exists(output_file):
                print(f"   - {output_file}")
        print("\n✨ Conversion complete!")
    else:
        print("\n⚠️  No files were converted successfully")
        sys.exit(1)

if __name__ == '__main__':
    main()
